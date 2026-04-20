"""Export OCR text to Word (.docx) with basic formatting detection."""

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _is_heading(line: str) -> bool:
    """Detect if a line is likely a heading (all caps, short, centered-looking)."""
    stripped = line.strip()
    if not stripped or len(stripped) > 100:
        return False
    # All caps or mostly caps
    alpha = [c for c in stripped if c.isalpha()]
    if alpha and sum(1 for c in alpha if c.isupper()) / len(alpha) > 0.7:
        return len(stripped) < 80
    return False


def _is_list_item(line: str) -> bool:
    """Detect bulleted or numbered list items."""
    stripped = line.strip()
    return bool(re.match(r'^[\-•·]\s', stripped) or re.match(r'^\d+[\.\)]\s', stripped))


def text_to_word(
    text: str,
    output_path: Path,
    font_name: str = "Times New Roman",
    font_size: int = 13,
    detect_format: bool = True,
) -> Path:
    """Convert OCR text to a formatted Word document.

    Args:
        text: OCR text content
        output_path: Where to save .docx
        font_name: Font for body text
        font_size: Font size in pt
        detect_format: Auto-detect headings, lists, etc.
    """
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.15

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            # Skip multiple blank lines, add one paragraph
            if i == 0 or lines[i - 1].strip():
                doc.add_paragraph("")
            i += 1
            continue

        if detect_format and _is_heading(stripped):
            p = doc.add_paragraph(stripped)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(font_size + 1)

        elif detect_format and _is_list_item(stripped):
            # Remove bullet/number prefix
            clean = re.sub(r'^[\-•·]\s*', '', stripped)
            clean = re.sub(r'^\d+[\.\)]\s*', '', clean)
            doc.add_paragraph(clean, style="List Bullet")

        else:
            doc.add_paragraph(stripped)

        i += 1

    doc.save(str(output_path))
    return output_path
